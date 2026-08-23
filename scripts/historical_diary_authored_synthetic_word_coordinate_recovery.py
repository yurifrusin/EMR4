"""Run two local authored-synthetic Word timeout/throughput recovery proofs."""

from __future__ import annotations

import json
from pathlib import Path
import subprocess
import sys
import time
from typing import Any

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from orchestration_harness import historical_diary_local_measured_privacy_probe as probe


BASE_ROOT = ROOT / "local_data/authored-synthetic-diary-word-coordinate-recovery"
RUN_ROOT = BASE_ROOT / "run-v1"
DOCUMENT_ROOT = RUN_ROOT / "documents"
MANIFEST_PATH = RUN_ROOT / "synthetic-binding-manifest.json"
CONTROL_PATH = RUN_ROOT / "owned-word-process-control.json"
PROGRESS_PATH = RUN_ROOT / "word-extraction-progress.json"
TIMEOUT_CONTROL_PATH = RUN_ROOT / "timeout-owned-word-process-control.json"
TIMEOUT_CLEANUP_PATH = RUN_ROOT / "timeout-cleanup-receipt.json"
THROUGHPUT_CLEANUP_PATH = RUN_ROOT / "throughput-cleanup-receipt.json"
EXTRACTOR_PATH = ROOT / "scripts/historical_diary_local_measured_privacy_probe.ps1"
TIMEOUT_FIXTURE_PATH = (
    ROOT / "scripts/historical_diary_authored_synthetic_word_timeout_fixture.ps1"
)
WORD_CLEANUP_PATH = ROOT / "scripts/historical_diary_owned_word_cleanup.ps1"

DOCUMENT_COUNT = 12
TABLE_ROWS = 7
TABLE_COLUMNS = 2
SEGMENTS_PER_CELL = 12
STORY_ANCHORS_PER_DOCUMENT = 24


class RecoveryError(RuntimeError):
    """A synthetic recovery proof failed without releasing a source value."""


def _word_process_ids() -> set[int]:
    completed = subprocess.run(
        [
            "powershell.exe",
            "-NoProfile",
            "-NonInteractive",
            "-Command",
            "@(Get-Process -Name WINWORD -ErrorAction SilentlyContinue | "
            "ForEach-Object { $_.Id }) | ConvertTo-Json -Compress",
        ],
        cwd=ROOT,
        capture_output=True,
        check=False,
        timeout=30,
        text=True,
        encoding="utf-8",
    )
    if completed.returncode != 0:
        raise RecoveryError("word_process_snapshot_failed")
    try:
        value = json.loads(completed.stdout or "[]")
    except json.JSONDecodeError as error:
        raise RecoveryError("word_process_snapshot_invalid") from error
    values = value if isinstance(value, list) else [value]
    if any(not isinstance(item, int) or item < 1 for item in values):
        raise RecoveryError("word_process_snapshot_invalid")
    return set(values)


def _prepare_root() -> None:
    if RUN_ROOT.exists():
        raise RecoveryError("synthetic_run_root_already_exists")
    DOCUMENT_ROOT.mkdir(parents=True)
    if RUN_ROOT.resolve().parent != BASE_ROOT.resolve():
        raise RecoveryError("synthetic_run_root_invalid")
    if "historical-diary-trove" in RUN_ROOT.as_posix():
        raise RecoveryError("historical_root_forbidden")


def _create_documents() -> tuple[Path, ...]:
    paths: list[Path] = []
    for document_index in range(DOCUMENT_COUNT):
        document = Document()
        for anchor_index in range(STORY_ANCHORS_PER_DOCUMENT):
            hour = 8 + (anchor_index // 6)
            minute = (anchor_index % 6) * 10
            document.add_paragraph(f"{hour:02}:{minute:02}")
        table = document.add_table(rows=TABLE_ROWS, cols=TABLE_COLUMNS)
        for cell_index, cell in enumerate(
            cell for row in table.rows for cell in row.cells
        ):
            cell.paragraphs[0].text = (
                f"SYNTHETIC SLOT {document_index:02}-{cell_index:02}-00"
            )
            for segment_index in range(1, SEGMENTS_PER_CELL):
                cell.add_paragraph(
                    f"SYNTHETIC SLOT {document_index:02}-{cell_index:02}-"
                    f"{segment_index:02}"
                )
        path = DOCUMENT_ROOT / f"synthetic-{document_index:02}.docx"
        document.save(path)
        paths.append(path)
    return tuple(paths)


def _write_manifest(paths: tuple[Path, ...]) -> None:
    if len(paths) != DOCUMENT_COUNT:
        raise RecoveryError("synthetic_document_count_invalid")
    payload = {
        "schema_version": "historical_diary.authored_synthetic_binding_manifest.v1",
        "files": [
            {
                "sequence_index": index,
                "observation_offset_seconds": index * 30,
                "absolute_path": str(path.resolve()),
            }
            for index, path in enumerate(paths)
        ],
    }
    MANIFEST_PATH.write_text(
        json.dumps(payload, indent=2) + "\n", encoding="utf-8", newline="\n"
    )


def _timeout_containment_proof() -> dict[str, Any]:
    baseline = _word_process_ids()
    command = [
        "powershell.exe",
        "-NoProfile",
        "-NonInteractive",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(TIMEOUT_FIXTURE_PATH),
        "-ControlPath",
        str(TIMEOUT_CONTROL_PATH),
    ]
    process = subprocess.Popen(
        command,
        cwd=ROOT,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    deadline = time.monotonic() + 30
    while time.monotonic() < deadline and not TIMEOUT_CONTROL_PATH.exists():
        if process.poll() is not None:
            raise RecoveryError("timeout_fixture_exited_before_control")
        time.sleep(0.1)
    if not TIMEOUT_CONTROL_PATH.exists():
        process.kill()
        process.wait(timeout=10)
        raise RecoveryError("timeout_fixture_control_not_ready")
    probe.OwnedWordProcessControl.model_validate_json(
        TIMEOUT_CONTROL_PATH.read_bytes()
    )
    try:
        process.communicate(timeout=2)
    except subprocess.TimeoutExpired:
        process.kill()
        process.communicate(timeout=10)
    else:
        raise RecoveryError("timeout_fixture_did_not_timeout")
    cleanup = probe.cleanup_owned_word_process(
        control_path=TIMEOUT_CONTROL_PATH,
        receipt_path=TIMEOUT_CLEANUP_PATH,
        cleanup_script_path=WORD_CLEANUP_PATH,
    )
    after = _word_process_ids()
    if (
        cleanup.reason_code != "owned_process_removed"
        or not cleanup.exact_owned_process_absent
        or after != baseline
    ):
        raise RecoveryError("timeout_containment_postcondition_failed")
    return {
        "typed_reason_code": "word_extractor_timeout",
        "owned_processes_removed": 1,
        "preexisting_process_set_preserved": True,
        "broad_process_name_stop_used": False,
    }


def _throughput_proof() -> dict[str, Any]:
    paths = _create_documents()
    _write_manifest(paths)
    baseline = _word_process_ids()
    command = [
        "powershell.exe",
        "-NoProfile",
        "-NonInteractive",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(EXTRACTOR_PATH),
        "-Manifest",
        str(MANIFEST_PATH),
        "-ControlPath",
        str(CONTROL_PATH),
        "-ProgressPath",
        str(PROGRESS_PATH),
        "-ExecutionProfile",
        "AuthoredSyntheticRecovery",
    ]
    completed = probe.run_owned_word_subprocess(
        command,
        timeout_seconds=300,
        control_path=CONTROL_PATH,
        cleanup_receipt_path=THROUGHPUT_CLEANUP_PATH,
        cleanup_script_path=WORD_CLEANUP_PATH,
    )
    try:
        extraction = probe.PrivateExtraction.model_validate_json(completed.stdout)
        progress = probe.WordExtractionProgress.model_validate_json(
            PROGRESS_PATH.read_bytes()
        )
    except (OSError, ValueError) as error:
        cleanup = probe.cleanup_owned_word_process(
            control_path=CONTROL_PATH,
            receipt_path=THROUGHPUT_CLEANUP_PATH,
            cleanup_script_path=WORD_CLEANUP_PATH,
        )
        if not cleanup.exact_owned_process_absent:
            raise RecoveryError("throughput_invalid_output_cleanup_failed") from error
        raise RecoveryError("throughput_output_invalid") from error
    after = _word_process_ids()
    if after != baseline:
        raise RecoveryError("throughput_preexisting_process_set_changed")
    if (
        extraction.status != "passed"
        or not extraction.word_cleanup_completed
        or len(extraction.snapshots) != DOCUMENT_COUNT
        or progress.stage != "cleanup"
        or progress.completed_document_count != DOCUMENT_COUNT
    ):
        raise RecoveryError("throughput_extraction_boundary_failed")
    return {
        "opened_document_count": len(extraction.snapshots),
        "parsed_document_count": sum(
            snapshot.error_code is None for snapshot in extraction.snapshots
        ),
        "table_cell_count": progress.table_cell_count,
        "structural_segment_count": progress.structural_segment_count,
        "coordinate_attempt_count": progress.coordinate_attempt_count,
        "explicit_story_anchor_count": progress.explicit_story_anchor_count,
        "elapsed_bucket": progress.elapsed_bucket,
        "coordinate_rate_floor_bucket": progress.coordinate_rate_floor_bucket,
        "word_cleanup_completed": extraction.word_cleanup_completed,
        "preexisting_process_set_preserved": True,
    }


def _cleanup_created_files() -> bool:
    success = True
    for path in (
        CONTROL_PATH,
        PROGRESS_PATH,
        TIMEOUT_CONTROL_PATH,
        TIMEOUT_CLEANUP_PATH,
        THROUGHPUT_CLEANUP_PATH,
        MANIFEST_PATH,
        Path(str(PROGRESS_PATH) + ".tmp"),
    ):
        try:
            path.unlink(missing_ok=True)
        except OSError:
            success = False
    if DOCUMENT_ROOT.exists():
        for path in sorted(DOCUMENT_ROOT.iterdir()):
            try:
                if path.is_file() and not path.is_symlink():
                    path.unlink()
                else:
                    success = False
            except OSError:
                success = False
        try:
            DOCUMENT_ROOT.rmdir()
        except OSError:
            success = False
    for path in (RUN_ROOT, BASE_ROOT):
        try:
            path.rmdir()
        except FileNotFoundError:
            pass
        except OSError:
            success = False
    return success


def main() -> int:
    timeout_result: dict[str, Any] | None = None
    throughput_result: dict[str, Any] | None = None
    decision = "blocked"
    reason_code = "internal_synthetic_recovery_failure"
    cleanup_completed = False
    try:
        _prepare_root()
        timeout_result = _timeout_containment_proof()
        throughput_result = _throughput_proof()
        if (
            throughput_result["opened_document_count"] == DOCUMENT_COUNT
            and throughput_result["parsed_document_count"] == DOCUMENT_COUNT
            and throughput_result["coordinate_attempt_count"] >= 2000
            and throughput_result["explicit_story_anchor_count"] >= 250
            and throughput_result["elapsed_bucket"]
            in {"under_30_seconds", "30_to_119_seconds", "120_to_299_seconds"}
            and throughput_result["coordinate_rate_floor_bucket"] != "not_available"
        ):
            decision = "passed"
            reason_code = "passed"
        else:
            decision = "revision_required"
            reason_code = "synthetic_throughput_floor_not_met"
    except (OSError, subprocess.SubprocessError, ValueError, RecoveryError, probe.ProbeError) as error:
        reason = str(error)
        reason_code = (
            reason
            if reason
            and len(reason) <= 100
            and all(character.islower() or character.isdigit() or character == "_" for character in reason)
            else "internal_synthetic_recovery_failure"
        )
    finally:
        if TIMEOUT_CONTROL_PATH.exists():
            probe.cleanup_owned_word_process(
                control_path=TIMEOUT_CONTROL_PATH,
                receipt_path=TIMEOUT_CLEANUP_PATH,
                cleanup_script_path=WORD_CLEANUP_PATH,
            )
        if CONTROL_PATH.exists():
            probe.cleanup_owned_word_process(
                control_path=CONTROL_PATH,
                receipt_path=THROUGHPUT_CLEANUP_PATH,
                cleanup_script_path=WORD_CLEANUP_PATH,
            )
        cleanup_completed = _cleanup_created_files()
        if not cleanup_completed:
            decision = "blocked"
            reason_code = "synthetic_cleanup_incomplete"
    result = {
        "schema_version": "historical_diary.authored_synthetic_word_coordinate_recovery.v1",
        "decision": decision,
        "reason_code": reason_code,
        "timeout_containment": timeout_result,
        "throughput": throughput_result,
        "cleanup_completed": cleanup_completed,
        "historical_archive_enumerations": 0,
        "historical_content_reads": 0,
        "source_value_emitted": False,
        "provider_network_model_calls": 0,
        "historical_derived_artifacts_created": 0,
        "first_use_gate_opened": False,
    }
    print(json.dumps(result, indent=2, sort_keys=True))
    return 0 if decision == "passed" else 2


if __name__ == "__main__":
    raise SystemExit(main())
