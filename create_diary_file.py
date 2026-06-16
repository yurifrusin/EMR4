"""
create_diary_file.py -- Generate a daily EMR4 Centaur Diary .docx

Increment 2: multi-column Word table with grey headers, time-slot stacks,
break rows, per-column footer (Messages / Phone Consultations), and
lifecycle-state colour rendering on booked patient names.

Appointment name formatting follows the patient encounter state machine:
  Booked     -> plain black      (receptionist entered, unconfirmed)
  Confirmed  -> ALL-CAPS, blue   (record located in system)
  Arrived    -> ALL-CAPS, blue   (chevron removed; same colour as Confirmed)
  InConsult  -> underlined       (doctor hit Ctrl+Alt+N)
  Completed  -> green            (doctor billed)
  others     -> plain black

The template is loaded from diary_template.json in the repo root. See that
file for the column/break/slot configuration format.

CLI usage:
    python create_diary_file.py
    python create_diary_file.py --date 16-06-2026 --out ./patient_files
    python create_diary_file.py --date 16-06-2026 --out ./patient_files --db
"""

import argparse
import json
import zipfile
from datetime import date, datetime, time, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ── Typography (match create_patient_file.py) ─────────────────────────────────
EMR4_BLUE    = RGBColor(0x00, 0x00, 0xFF)
EMR4_GREEN   = RGBColor(0x00, 0xB0, 0x50)
BODY_FONT    = "Century Schoolbook"
HEADING_FONT = "Garamond"
BODY_PT      = 11
HEADING_PT   = 12
HEADER_GREY  = "E6E6E6"
BREAK_GREY   = "F2F2F2"

TEMPLATE_PATH = Path(__file__).parent / "diary_template.json"


# ── Custom XML ────────────────────────────────────────────────────────────────
_CUSTOM_XML_NS = "http://emr4.com/ns/document"
_CUSTOM_XML_CONTENT = (
    '<?xml version="1.0" encoding="UTF-8"?>\n'
    '<emr4:root xmlns:emr4="{ns}">\n'
    '  <emr4:document-type>diary</emr4:document-type>\n'
    '  <emr4:version>1.0</emr4:version>\n'
    '</emr4:root>'
).format(ns=_CUSTOM_XML_NS)

_CUSTOM_XML_RELS_ENTRY = (
    '<Relationship Id="rIdEMR4Custom" '
    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml" '
    'Target="../customXml/item1.xml"/>'
)

_CUSTOM_XML_PROPS = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<ds:datastoreItem ds:itemID="{EMR4-0002-0000-0000-000000000001}"\n'
    '  xmlns:ds="http://schemas.openxmlformats.org/officeDocument/2006/customXml">\n'
    '  <ds:schemaRefs/>\n'
    '</ds:datastoreItem>'
)


# ── OOXML helpers (ordered insertion — Word Online is strict about schema order) ─

def _insert_element_before(parent, new_child, before_tag):
    """Insert new_child immediately before the first occurrence of before_tag.
    Falls back to append if the tag is absent.  Required because Word Online
    enforces CT_PPr / CT_TcPr child-element order; Desktop silently tolerates it."""
    ref = parent.find(qn(before_tag))
    if ref is not None:
        ref.addprevious(new_child)
    else:
        parent.append(new_child)


def _shade_paragraph(para, fill_hex):
    """Apply a solid background to a paragraph's <w:pPr>.
    w:shd must precede w:spacing / w:jc in CT_PPr (Word Online strictness)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    _insert_element_before(pPr, shd, "w:spacing")


def _shade_cell(cell, fill_hex):
    """Apply a solid background to a table cell via <w:tcPr><w:shd>.
    In CT_TcPr, w:shd must appear after w:tcBorders and before w:noWrap."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    _insert_element_before(tcPr, shd, "w:noWrap")


def _apply_template_styles(doc):
    """Set Normal and Heading 1 fonts to match the EMR4 patient-file template."""
    normal = doc.styles["Normal"]
    nf = normal.font
    nf.name = BODY_FONT
    nf.size = Pt(BODY_PT)

    h1 = doc.styles["Heading 1"]
    h1f = h1.font
    h1f.name = HEADING_FONT
    h1f.size = Pt(HEADING_PT)
    h1f.bold = True
    h1f.color.rgb = EMR4_BLUE


def _set_run(run, bold=False, underline=False, color=None,
             size=BODY_PT, font_name=BODY_FONT):
    """Apply standard font properties to a run."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color


# ── Custom XML injection (identical fix applied to both generators) ───────────

def _inject_custom_xml(doc_path):
    """Inject the EMR4 Custom XML Part (document-type=diary) into the .docx zip.

    Overwrites only the body of customXml/item1.xml; leaves all OPC packaging
    intact (no duplicate Override entries — the corruption bug from b160dd0).
    Defensive fallback if python-docx ever ships without the customXml slot.
    """
    import shutil

    _ITEM_RELS = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rIdProps1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXmlProps" '
        'Target="itemProps1.xml"/>'
        '</Relationships>'
    )

    tmp = doc_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(doc_path, "r") as zin:
        names = set(zin.namelist())
        has_slot = "customXml/item1.xml" in names

        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "customXml/item1.xml":
                    # Replace the Bibliography stub body with our marker —
                    # all other OPC packaging (itemProps, rels, content-types) unchanged.
                    zout.writestr(item, _CUSTOM_XML_CONTENT.encode("utf-8"))
                elif not has_slot and item.filename == "[Content_Types].xml":
                    # Fallback path only: add the two missing Overrides.
                    text = zin.read(item.filename).decode("utf-8")
                    if "customXml/item1.xml" not in text:
                        entry = '<Override PartName="/customXml/item1.xml" ContentType="application/xml"/>'
                        props = '<Override PartName="/customXml/itemProps1.xml" ContentType="application/vnd.openxmlformats-officedocument.customXmlProperties+xml"/>'
                        text = text.replace("</Types>", entry + props + "</Types>")
                    zout.writestr(item, text.encode("utf-8"))
                elif not has_slot and item.filename == "word/_rels/document.xml.rels":
                    # Fallback path only: add the missing relationship.
                    rels = zin.read(item.filename).decode("utf-8")
                    if "rIdEMR4Custom" not in rels:
                        rels = rels.replace(
                            "</Relationships>",
                            "  " + _CUSTOM_XML_RELS_ENTRY + "\n</Relationships>",
                        )
                    zout.writestr(item, rels.encode("utf-8"))
                else:
                    zout.writestr(item, zin.read(item.filename))

            if not has_slot:
                # Fallback: create all three slot files from scratch.
                zout.writestr("customXml/item1.xml", _CUSTOM_XML_CONTENT.encode("utf-8"))
                zout.writestr("customXml/itemProps1.xml", _CUSTOM_XML_PROPS.encode("utf-8"))
                zout.writestr("customXml/_rels/item1.xml.rels", _ITEM_RELS.encode("utf-8"))

    shutil.move(str(tmp), str(doc_path))


# ── Template loading ──────────────────────────────────────────────────────────

def _load_template():
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Diary template not found: {TEMPLATE_PATH}\n"
            "Create diary_template.json in the repo root."
        )
    with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


# ── Slot generation ───────────────────────────────────────────────────────────

def _generate_slots(template):
    """Return an ordered list of (time_str, kind, label) tuples.

    kind='slot'  -> time_str is e.g. '09:00', label is None
    kind='break' -> time_str is '',           label is the display text
    """
    defaults = template["slot_defaults"]
    start_t  = datetime.strptime(defaults["start"], "%H:%M").time()
    end_t    = datetime.strptime(defaults["end"],   "%H:%M").time()
    interval = timedelta(minutes=defaults["interval_minutes"])

    breaks = []
    for b in template.get("breaks", []):
        breaks.append({
            "label":  b["label"],
            "from_t": datetime.strptime(b["from"], "%H:%M").time(),
            "to_t":   datetime.strptime(b["to"],   "%H:%M").time(),
        })

    # Use a fixed reference date for time arithmetic
    REF = date(2000, 1, 1)
    current = datetime.combine(REF, start_t)

    slots = []
    inserted_breaks = set()

    while current.time() < end_t:
        t = current.time()

        # Check whether this time falls inside a break window
        in_break = None
        for i, b in enumerate(breaks):
            if b["from_t"] <= t < b["to_t"]:
                in_break = i
                break

        if in_break is not None:
            # Emit the break label once, then skip all times in that window
            if in_break not in inserted_breaks:
                inserted_breaks.add(in_break)
                slots.append(("", "break", breaks[in_break]["label"]))
        else:
            slots.append((current.strftime("%H:%M"), "slot", None))

        current += interval

    return slots


# ── Appointment lookup ────────────────────────────────────────────────────────

def _build_appt_lookup(template, appointments):
    """Map appointments to columns by practitioner AHPRA number.

    Returns:
        {col_idx: {time_str: (patient_full_name, status_str)}}

    Multiple bookings at the same time (family/multi-name slots) are joined
    with ' / ' — reception can edit further in the Word doc.
    """
    ahpra_to_col = {}
    for i, col in enumerate(template["columns"]):
        ahpra = col.get("practitioner_ahpra")
        if ahpra:
            ahpra_to_col[ahpra] = i

    lookup = {}
    for appt in appointments:
        ahpra = appt.get("practitioner_ahpra")
        col_idx = ahpra_to_col.get(ahpra)
        if col_idx is None:
            continue

        start = appt["start_time"]
        time_key = start.strftime("%H:%M") if isinstance(start, datetime) else start
        name   = "{} {}".format(appt["patient_first_name"], appt["patient_last_name"])
        status = appt.get("status", "Booked")

        if col_idx not in lookup:
            lookup[col_idx] = {}

        if time_key in lookup[col_idx]:
            # Family booking: append name (status from the first booking)
            prev_name, prev_status = lookup[col_idx][time_key]
            lookup[col_idx][time_key] = ("{} / {}".format(prev_name, name), prev_status)
        else:
            lookup[col_idx][time_key] = (name, status)

    return lookup


# ── Cell builders ─────────────────────────────────────────────────────────────

def _fill_header_cell(cell, col_cfg, diary_date):
    """Grey header cell: bold assignment name, room label, centred."""
    _shade_cell(cell, HEADER_GREY)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after  = Pt(3)

    name_run = para.add_run(col_cfg["assignment"])
    _set_run(name_run, bold=True, font_name=HEADING_FONT, size=HEADING_PT)

    room_run = para.add_run("\n{}".format(col_cfg["room_label"]))
    _set_run(room_run, size=BODY_PT - 1)


def _add_slot_para(cell, time_str, patient_name, status):
    """Add one time-slot paragraph to a body cell.

    Empty slot  → '09:00  »'
    Booked      → '09:00  Margaret Thompson'      (plain black)
    Confirmed   → '09:00  MARGARET THOMPSON'      (ALL-CAPS, blue)
    Arrived     → '09:00  MARGARET THOMPSON'      (ALL-CAPS, blue — chevron removed)
    InConsult   → '09:00  Margaret Thompson'      (underlined)
    Completed   → '09:00  Margaret Thompson'      (green)
    """
    para = cell.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)

    if patient_name:
        # Time label
        time_run = para.add_run("{}  ".format(time_str))
        _set_run(time_run)

        # Name formatted by lifecycle status
        s = status or "Booked"
        if s in ("Confirmed", "Arrived"):
            name_run = para.add_run(patient_name.upper())
            _set_run(name_run, bold=True, color=EMR4_BLUE)
        elif s == "InConsult":
            name_run = para.add_run(patient_name)
            _set_run(name_run, underline=True)
        elif s == "Completed":
            name_run = para.add_run(patient_name)
            _set_run(name_run, color=EMR4_GREEN)
        else:
            # Booked, Cancelled, NoShow, DNA → plain black
            name_run = para.add_run(patient_name)
            _set_run(name_run)
    else:
        # Empty bookable slot
        slot_run = para.add_run("{}  »".format(time_str))  # » U+00BB
        _set_run(slot_run)


def _add_break_para(cell, label):
    """Bold break-row paragraph with a light grey background."""
    para = cell.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    _shade_paragraph(para, BREAK_GREY)
    run = para.add_run(label)
    _set_run(run, bold=True)


def _fill_body_cell(cell, col_cfg, slots, appt_col, footer):
    """Fill the body cell with time slots, break rows, and the footer.

    The initial empty paragraph left by python-docx stays as a zero-height
    top spacer (simpler than deleting it and safe for the OPC structure).
    """
    # Minimise the initial paragraph's spacing
    init = cell.paragraphs[0]
    init.paragraph_format.space_before = Pt(0)
    init.paragraph_format.space_after  = Pt(0)

    # Apply optional whole-column tint (e.g. yellow for nurse column).
    # Break-row paragraph shading will override this for those rows.
    tint = col_cfg.get("tint")
    if tint:
        _shade_cell(cell, tint)

    # Slot + break rows
    for time_str, kind, label in slots:
        if kind == "break":
            _add_break_para(cell, label)
        else:
            booked = appt_col.get(time_str)
            if booked:
                patient_name, status = booked
                _add_slot_para(cell, time_str, patient_name, status)
            else:
                _add_slot_para(cell, time_str, None, None)

    # Footer (Messages: / Phone Consultations:)
    if footer:
        sep = cell.add_paragraph()
        sep.paragraph_format.space_before = Pt(6)
        sep.paragraph_format.space_after  = Pt(0)
        _set_run(sep.add_run("─" * 24), size=8)  # ─

        for label in footer:
            lp = cell.add_paragraph()
            lp.paragraph_format.space_before = Pt(4)
            lp.paragraph_format.space_after  = Pt(10)
            _set_run(lp.add_run(label), bold=True)


# ── Main diary generator ──────────────────────────────────────────────────────

def create_diary_docx(diary_date, output_dir=None, appointments=None):
    """Generate Diary_DD-MM-YYYY.docx for the given date.

    Args:
        diary_date:   datetime.date for the diary page.
        output_dir:   pathlib.Path; defaults to current directory.
        appointments: list of dicts with keys:
                        start_time (datetime), patient_first_name,
                        patient_last_name, status (AppointmentStatus value
                        string), practitioner_ahpra.
    Returns the Path of the written file.
    """
    if output_dir is None:
        output_dir = Path(".")

    template   = _load_template()
    slots      = _generate_slots(template)
    appt_lookup = _build_appt_lookup(template, appointments or [])

    doc = Document()
    _apply_template_styles(doc)

    # Date heading (Heading 1, centred, shaded grey)
    day_label = diary_date.strftime("%A %d %B %Y").replace(" 0", " ", 1)
    h = doc.add_heading("Diary  {}".format(day_label), level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade_paragraph(h, HEADER_GREY)

    doc.add_paragraph()  # breathing room before the table

    # Multi-column table: 2 rows (header + body) × N columns
    columns = template["columns"]
    n = len(columns)
    table = doc.add_table(rows=2, cols=n)
    table.style = "Table Grid"

    # Row 0 — grey header cells
    for i, col_cfg in enumerate(columns):
        _fill_header_cell(table.cell(0, i), col_cfg, diary_date)

    # Row 1 — body cells (time slots + breaks + footer)
    footer = template.get("footer", [])
    for i, col_cfg in enumerate(columns):
        _fill_body_cell(
            table.cell(1, i),
            col_cfg,
            slots,
            appt_lookup.get(i, {}),
            footer,
        )

    # Write to disk
    date_str = diary_date.strftime("%d-%m-%Y")
    filename = "Diary_{}.docx".format(date_str)
    out_path = Path(output_dir) / filename
    doc.save(str(out_path))

    # Inject Custom XML Part (document-type=diary)
    _inject_custom_xml(out_path)

    return out_path


# ── DB loader (used by --db CLI flag) ────────────────────────────────────────

def _load_appointments_from_db(diary_date):
    """Pull appointments for diary_date from the local dev DB.

    Returns a list of dicts in the format expected by create_diary_docx's
    appointments parameter.  Imports app models lazily so the script can
    also be used without the venv active (just without --db).
    """
    import sys
    if "" not in sys.path:
        sys.path.insert(0, "")

    from app.database import SessionLocal
    from app.models.appointments import Appointment

    db = SessionLocal()
    try:
        start_dt = datetime.combine(diary_date, time(0, 0))
        end_dt   = datetime.combine(diary_date, time(23, 59, 59))

        rows = (
            db.query(Appointment)
            .filter(
                Appointment.start_time >= start_dt,
                Appointment.start_time <= end_dt,
            )
            .all()
        )

        result = []
        for appt in rows:
            pt = appt.patient
            pr = appt.practitioner
            result.append({
                "start_time":         appt.start_time,
                "patient_first_name": pt.first_name,
                "patient_last_name":  pt.last_name,
                "status":             appt.status.value,
                "practitioner_ahpra": pr.ahpra_number,
            })
        return result
    finally:
        db.close()


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate an EMR4 diary .docx for a given date."
    )
    parser.add_argument(
        "--date", default=None,
        help="Date as DD-MM-YYYY (default: today)"
    )
    parser.add_argument(
        "--out", default=".",
        help="Output directory (default: .)"
    )
    parser.add_argument(
        "--db", action="store_true",
        help="Pull appointments from the local dev DB (requires venv + running Postgres)"
    )
    args = parser.parse_args()

    if args.date:
        diary_date = datetime.strptime(args.date, "%d-%m-%Y").date()
    else:
        diary_date = date.today()

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    appts = _load_appointments_from_db(diary_date) if args.db else []
    path  = create_diary_docx(diary_date, output_dir=out_dir, appointments=appts)
    print("Created: {}".format(path))
    if appts:
        print("  Populated with {} appointment(s) from DB".format(len(appts)))
