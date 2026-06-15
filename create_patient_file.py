"""
create_patient_file.py — Generate a new EMR4 Centaur patient .docx

The core function `create_patient_docx()` is the integration point for the future
New Patient userform.  Import it from a FastAPI endpoint like so:

    from create_patient_file import create_patient_docx, PatientData
    path = create_patient_docx(patient, output_dir=Path("/mnt/sharepoint/patients"))

CLI usage:
    python create_patient_file.py
    python create_patient_file.py --first Billy --last Frusin --dob 15-10-2015 \\
        --sex Male --address "15 Rose St Blacktown NSW 2148" \\
        --phone "0412 345 678" --medicare "1234567891"
"""

import argparse
import re
import shutil
import zipfile
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Section headings — Dr Shera structure ─────────────────────────────────────
# (heading text, content-control tag). Order, text, AND tags must match
# PROTECTED_SECTIONS in taskpane.js. "Contemporaneous Notes" is the AI anchor.
SECTION_HEADINGS = [
    ("Care Plans, Health Assessments and Recalls", "emr4-section-care-plans"),
    ("Family History",                             "emr4-section-family-history"),
    ("Medical History",                            "emr4-section-medical-history"),
    ("Social History",                             "emr4-section-social-history"),
    ("Current Drugs",                              "emr4-section-current-drugs"),
    ("Drug Reactions",                             "emr4-section-drug-reactions"),
    ("Contemporaneous Notes",                      "emr4-section-cn"),
    ("Vaccinations",                               "emr4-section-vaccinations"),
    ("Specialist Reports",                         "emr4-section-specialist-reports"),
    ("Diagnostic Imaging",                         "emr4-section-imaging"),
    ("Pathology Results",                          "emr4-section-pathology"),
    ("ECG Records",                                "emr4-section-ecg"),
    ("Prescription Records",                       "emr4-section-prescription-records"),
    ("Correspondence",                             "emr4-section-correspondence"),
    ("Management Articles",                        "emr4-section-management-articles"),
]

# Typography — must match the Margaret Thompson reference file (the template).
EMR4_BLUE    = RGBColor(0x00, 0x00, 0xFF)   # pure blue 0000FF
BODY_FONT    = "Century Schoolbook"          # Normal style + demographics
HEADING_FONT = "Garamond"                    # Heading 1 section titles
BODY_PT      = 11
HEADING_PT   = 12
HEADER_GREY  = "E6E6E6"   # grey demographics band — matches MT template
HEADER_LINE_SPACING = 1.15

_CUSTOM_XML_NS = "http://emr4.com/ns/document"
_CUSTOM_XML_CONTENT = f"""<?xml version="1.0" encoding="UTF-8"?>
<emr4:root xmlns:emr4="{_CUSTOM_XML_NS}">
  <emr4:document-type>patient</emr4:document-type>
  <emr4:version>1.0</emr4:version>
</emr4:root>"""

_CUSTOM_XML_RELS_ENTRY = (
    '<Relationship Id="rIdEMR4Custom" '
    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml" '
    'Target="../customXml/item1.xml"/>'
)

_CUSTOM_XML_PROPS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<ds:datastoreItem ds:itemID="{EMR4-0001-0000-0000-000000000001}"
  xmlns:ds="http://schemas.openxmlformats.org/officeDocument/2006/customXml">
  <ds:schemaRefs/>
</ds:datastoreItem>"""


@dataclass
class PatientData:
    """All fields the New Patient userform will collect."""
    first_name:      str
    last_name:       str
    date_of_birth:   date
    sex:             str            # "Male" | "Female" | "Other"
    address:         str = ""
    phone:           str = ""
    medicare_number: str = ""


# ── Helpers ───────────────────────────────────────────────────────────────────

def _clean(s: str) -> str:
    """Collapse internal whitespace runs to single spaces and trim. Applied to
    user-supplied field values so e.g. a double-spaced address doesn't trigger
    Word's grammar-error underline. The fixed layout separators between fields
    (name / dob / age, the Phone↔Medicare gap) are literals in the f-strings and
    are NOT affected by this."""
    return re.sub(r"\s+", " ", s or "").strip()


def _age(dob: date) -> int:
    today = date.today()
    return today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))


def _shade_paragraph(paragraph, fill_hex: str) -> None:
    """Apply solid background shading to a whole paragraph (the grey demographics
    band). Matches the Margaret Thompson template, which uses paragraph shading —
    NOT a table — so it renders reliably in Word Online as well as desktop.

    The <w:shd> element MUST be placed before <w:spacing>/<w:jc> per the CT_PPr
    schema order. Word desktop tolerates a misplaced shd, but Word Online silently
    ignores it (background fails to render), so we insert it at the correct spot."""
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:shd"))
    if existing is not None:
        pPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    # Insert before the first element that legally follows shd in CT_PPr order.
    pPr.insert_element_before(
        shd,
        "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
        "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
        "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
        "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
        "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
        "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
    )


def _apply_template_styles(doc) -> None:
    """Set Normal and Heading 1 styles to match the Margaret Thompson template."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)

    h1 = doc.styles["Heading 1"]
    h1.font.name       = HEADING_FONT
    h1.font.size       = Pt(HEADING_PT)
    h1.font.bold       = True
    h1.font.color.rgb  = EMR4_BLUE
    h1.paragraph_format.space_before = Pt(10)  # 200 twips
    h1.paragraph_format.space_after  = Pt(2)   # 40 twips


def _wrap_heading_in_locked_cc(paragraph, tag: str, alias: str, cc_id: int) -> None:
    """Wrap a heading paragraph in a block-level content control that is locked
    against deletion AND content editing (w:lock = sdtContentLocked). This is the
    OOXML equivalent of Office.js cannotDelete=true + cannotEdit=true, and matches
    the tags repairDocumentStructure() looks for in taskpane.js."""
    p = paragraph._p
    parent = p.getparent()
    idx = list(parent).index(p)

    sdt        = OxmlElement("w:sdt")
    sdtPr      = OxmlElement("w:sdtPr")
    sdtContent = OxmlElement("w:sdtContent")

    alias_el = OxmlElement("w:alias"); alias_el.set(qn("w:val"), alias)
    tag_el   = OxmlElement("w:tag");   tag_el.set(qn("w:val"), tag)
    id_el    = OxmlElement("w:id");    id_el.set(qn("w:val"), str(cc_id))
    lock_el  = OxmlElement("w:lock");  lock_el.set(qn("w:val"), "sdtContentLocked")

    for el in (alias_el, tag_el, id_el, lock_el):
        sdtPr.append(el)

    sdt.append(sdtPr)
    sdt.append(sdtContent)

    parent.insert(idx, sdt)   # put the sdt where the paragraph was
    parent.remove(p)          # detach the paragraph
    sdtContent.append(p)      # re-home it inside the content control


def _inject_custom_xml(docx_path: Path) -> None:
    """Injects emr4:document-type=patient Custom XML Part so the taskpane
    auto-detects this as a patient file and activates patient mode."""
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin:
        existing = set(zin.namelist())
        rels_xml = zin.read("word/_rels/document.xml.rels").decode("utf-8")
        if "rIdEMR4Custom" not in rels_xml:
            rels_xml = rels_xml.replace(
                "</Relationships>",
                f"  {_CUSTOM_XML_RELS_ENTRY}\n</Relationships>",
            )
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = (
                    rels_xml.encode()
                    if item.filename == "word/_rels/document.xml.rels"
                    else zin.read(item.filename)
                )
                zout.writestr(item, data)
            if "customXml/item1.xml" not in existing:
                zout.writestr("customXml/item1.xml", _CUSTOM_XML_CONTENT)
            if "customXml/itemProps1.xml" not in existing:
                zout.writestr("customXml/itemProps1.xml", _CUSTOM_XML_PROPS)
    shutil.move(str(tmp), str(docx_path))


# ── Core function ─────────────────────────────────────────────────────────────

def create_patient_docx(patient: PatientData, output_dir: Path = Path(".")) -> Path:
    """
    Generate an EMR4 patient .docx and return its path.

    This is the function a FastAPI /patients/{id}/create-file endpoint should call.
    The file is named  FIRSTNAME LASTNAME DD-MM-YYYY.docx  so the taskpane
    auto-detect logic (autoDetectPatient) can identify the patient from the filename.
    """
    doc = Document()
    _apply_template_styles(doc)

    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Normalise user-supplied fields (collapse stray double spaces, etc.)
    first    = _clean(patient.first_name)
    last     = _clean(patient.last_name)
    sex      = _clean(patient.sex)
    address  = _clean(patient.address)
    phone    = _clean(patient.phone)
    medicare = _clean(patient.medicare_number)

    dob_str = patient.date_of_birth.strftime("%d-%m-%Y")
    age     = _age(patient.date_of_birth)
    name_uc = f"{first.upper()} {last.upper()}"

    # ── Demographics header ───────────────────────────────────────────────────
    # Grey shaded band made of three centred, shaded paragraphs — exactly how the
    # Margaret Thompson template does it (paragraph shading, no table). Tables with
    # hidden borders render inconsistently in Word Online; shaded paragraphs don't.
    def _header_line(text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        pf.line_spacing = HEADER_LINE_SPACING   # 1.15 lines, tight band
        _shade_paragraph(p, HEADER_GREY)
        run = p.add_run(text)
        run.font.name      = BODY_FONT
        run.font.bold      = True
        run.font.color.rgb = EMR4_BLUE
        run.font.size      = Pt(BODY_PT)

    # Always render all three demographic lines so the field structure is present
    # even when a value is blank — the receptionist/GP can fill gaps in-document,
    # and a userform-generated file shows the same layout as the template.
    _header_line(f"{name_uc}   dob {dob_str}   {age} years old   {sex}")
    _header_line(address)
    _header_line(f"Phone: {phone}        Medicare: {medicare}")

    doc.add_paragraph()  # spacer between header and first section

    # ── Section headings (each wrapped in a locked content control) ────────────
    for n, (heading_text, tag) in enumerate(SECTION_HEADINGS):
        h = doc.add_heading(heading_text, level=1)
        for run in h.runs:
            run.font.name      = HEADING_FONT
            run.font.color.rgb = EMR4_BLUE
        # Protect the heading: cannot be deleted or reformatted in Word
        _wrap_heading_in_locked_cc(h, tag=tag, alias=heading_text, cc_id=900000 + n)
        doc.add_paragraph()  # blank line under each heading (outside the CC) to write into

    # ── Save & inject Custom XML Part ────────────────────────────────────────
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{name_uc} {dob_str}.docx"
    out_path = output_dir / filename
    doc.save(str(out_path))
    _inject_custom_xml(out_path)
    return out_path


# ── CLI ───────────────────────────────────────────────────────────────────────

def _prompt(label: str, default: str = "") -> str:
    suffix = f" [{default}]" if default else ""
    return input(f"  {label}{suffix}: ").strip() or default


def _parse_dob(s: str) -> date:
    for fmt in ("%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    raise ValueError(f"Unrecognised date: {s!r}  — use DD-MM-YYYY")


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Create an EMR4 patient .docx")
    ap.add_argument("--first",    metavar="NAME",       help="First name")
    ap.add_argument("--last",     metavar="NAME",       help="Last name")
    ap.add_argument("--dob",      metavar="DD-MM-YYYY", help="Date of birth")
    ap.add_argument("--sex",      choices=["Male", "Female", "Other"])
    ap.add_argument("--address",  default=None,          help="Full address")
    ap.add_argument("--phone",    default=None,          help="Phone number")
    ap.add_argument("--medicare", default=None,          help="Medicare number")
    ap.add_argument("--out",      default=".",          metavar="DIR",
                    help="Output directory (default: current directory)")
    args = ap.parse_args()

    print("\nEMR4 — New Patient File Generator")
    print("-" * 36)

    # Required fields — prompt if not supplied via CLI
    first   = args.first or _prompt("First name")
    last    = args.last  or _prompt("Last name")
    dob_raw = args.dob   or _prompt("Date of birth (DD-MM-YYYY)")
    sex     = args.sex   or _prompt("Sex", "Female")

    # Optional fields — only prompt when running fully interactively (no CLI args at all)
    cli_mode = bool(args.first or args.last or args.dob or args.sex)
    address  = args.address  or ("" if cli_mode else _prompt("Address (optional)"))
    phone    = args.phone    or ("" if cli_mode else _prompt("Phone (optional)"))
    medicare = args.medicare or ("" if cli_mode else _prompt("Medicare number (optional)"))

    patient = PatientData(
        first_name      = first.strip(),
        last_name       = last.strip(),
        date_of_birth   = _parse_dob(dob_raw),
        sex             = sex,
        address         = address,
        phone           = phone,
        medicare_number = medicare,
    )

    out = create_patient_docx(patient, output_dir=Path(args.out))
    print(f"\n[OK] {out}")
